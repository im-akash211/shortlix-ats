"""
Text extraction service for PDF and DOCX resume files.

Raises TextExtractionError for all extraction failures so callers
can map it to a specific ingestion status without handling raw exceptions.
"""

import logging

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Raised when text cannot be extracted from a resume file."""


def extract_text_from_pdf(file_obj) -> str:
    """
    Extract plain text from a PDF file object.

    Uses pdfplumber which handles multi-column layouts better than
    pdfminer.six alone. Returns empty string for scanned pages (no text layer).
    """
    try:
        with pdfplumber.open(file_obj) as pdf:
            page_texts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    page_texts.append(text)
        return "\n".join(page_texts).strip()
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", exc)
        raise TextExtractionError(f"Could not read PDF file: {exc}") from exc


def extract_text_from_docx(file_obj) -> str:
    """Extract plain text from a DOCX file object."""
    try:
        doc = Document(file_obj)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        raise TextExtractionError(f"Could not read DOCX file: {exc}") from exc


def extract_text(file_obj, file_type: str) -> str:
    """
    Dispatch to the correct extractor based on file_type ('pdf' or 'docx').
    Returns the extracted raw text string.
    """
    normalized = file_type.lower().lstrip(".")
    if normalized == "pdf":
        return extract_text_from_pdf(file_obj)
    elif normalized == "docx":
        return extract_text_from_docx(file_obj)
    else:
        raise TextExtractionError(f"Unsupported file type: {file_type}")
